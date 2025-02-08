import json
import os
import io
from flask import Flask, request, render_template, redirect, url_for, flash
from openai import OpenAI
from dotenv import load_dotenv
from werkzeug.utils import secure_filename

# =============================================================================
# Load Environment Variables
# =============================================================================
load_dotenv()

# =============================================================================
# Initialize Flask App
# =============================================================================
app = Flask(__name__)
app.secret_key = 'your_secret_key'  # Needed for flashing error messages

# =============================================================================
# Initialize OpenAI Client
# =============================================================================
client = OpenAI(
    api_key=os.environ.get("OPENAI_API_KEY"),  # This is the default and can be omitted
)
if not client.api_key:
    raise ValueError("The OpenAI API key is not set. Please set it in the .env file.")

# =============================================================================
# Load GPT Prompt Configuration
# =============================================================================
with open('prompts.json', 'r') as file:
    prompts = json.load(file)

generate_document_prompt = prompts.get('generate_document_from_template_prompt')
if not generate_document_prompt:
    raise ValueError("The generate_document_from_template_prompt is not defined in prompts.json.")

summarize_document_prompt = prompts.get('summarize_document_prompt')
if not generate_document_prompt:
    raise ValueError("The summaize_document_from_template_prompt is not defined in prompts.json.")

# =============================================================================
# Utility Function: Read Uploaded File
# =============================================================================
def read_uploaded_file(uploaded_file):
    """
    Read the uploaded file and extract text depending on its file type.
    Supports .txt, .md, .docx, and .pdf files.
    """
    filename = secure_filename(uploaded_file.filename)
    ext = os.path.splitext(filename)[1].lower()
    
    # For text-based files
    if ext in ['.txt', '.md']:
        return uploaded_file.read().decode('utf-8')
    
    # For DOCX files
    elif ext == '.docx':
        from docx import Document
        doc = Document(uploaded_file)
        full_text = []
        for para in doc.paragraphs:
            full_text.append(para.text)
        return "\n".join(full_text)
    
    # For PDF files
    elif ext == '.pdf':
        import PyPDF2
        uploaded_file.seek(0)
        reader = PyPDF2.PdfReader(uploaded_file)
        pages_text = []
        for page in reader.pages:
            text = page.extract_text()
            if text:
                pages_text.append(text)
        return "\n".join(pages_text)
    
    # Fallback to reading as text
    else:
        return uploaded_file.read().decode('utf-8')

# =============================================================================
# Function: Process Additional Context Files
# =============================================================================
def process_context_files(context_files, query_text):
    """
    Process additional context files uploaded by the user.
    Reads and extracts text from the files, splits the text into chunks,
    converts chunks into Document objects, indexes them using a vector store,
    and retrieves relevant chunks via similarity search using the query_text.

    Args:
        context_files (list): List of file objects from the request.
        query_text (str): The text to use as a query for similarity search (e.g., template text).

    Returns:
        list: A list of Document objects that are most relevant or an empty list if no context is provided.
    """
    # If no context files are provided, return an empty list.
    if not context_files or context_files[0].filename == "":
        return []
    
    # Read and extract text from each context file.
    context_texts = []
    for file in context_files:
        context_texts.append(read_uploaded_file(file))
    
    # Split the text into chunks and convert each chunk into a Document.
    from langchain.text_splitter import RecursiveCharacterTextSplitter
    from langchain.docstore.document import Document
    all_context_docs = []
    text_splitter = RecursiveCharacterTextSplitter(chunk_size=1000, chunk_overlap=200)
    for doc_text in context_texts:
        chunks = text_splitter.split_text(doc_text)
        for chunk in chunks:
            all_context_docs.append(Document(page_content=chunk))
    
    # Build a vector store from the Document objects and perform similarity search.
    from langchain.embeddings import OpenAIEmbeddings
    from langchain.vectorstores import FAISS
    embeddings = OpenAIEmbeddings()
    vector_store = FAISS.from_documents(all_context_docs, embeddings)
    retrieved_docs = vector_store.similarity_search(query_text, k=5)
    return retrieved_docs

# =============================================================================
# Function: Summarize Long Document
# =============================================================================
def summarize_long_document(document_text):
    from langchain.text_splitter import RecursiveCharacterTextSplitter
    from langchain.docstore.document import Document
    
    # Split the long document into smaller chunks.
    text_splitter = RecursiveCharacterTextSplitter(chunk_size=1000, chunk_overlap=200)
    texts = text_splitter.split_text(document_text)
    
    # Convert each chunk into a Document object.
    docs = [Document(page_content=t) for t in texts]
    
    # Use the chat-based LLM wrapper for LangChain.
    from langchain.chat_models import ChatOpenAI
    llm = ChatOpenAI(temperature=0.5, model="gpt-4o")
    
    # Load a summarization chain (map_reduce is a good choice for long documents).
    from langchain.chains.summarize import load_summarize_chain
    chain = load_summarize_chain(llm, chain_type="map_reduce")
    
    # Process each document chunk individually.
    partial_summaries = []
    for doc in docs:
        summary = chain.run([doc])
        partial_summaries.append(summary)
    
    # Combine the partial summaries into a single summary.
    long_summary = "\n".join(partial_summaries)
    return long_summary


# =============================================================================
# Function: Summarize Document (with LangChain)
# =============================================================================
def summarize_document(document_text, template_text, LONG_DOC_THRESHOLD=3000):
    """
    Summarize the given document text with reference to the template.
    If the document is long, use LangChain's text splitter and summarization chain to
    generate a concise summary; otherwise, use the standard summarization prompt.
    
    Args:
        document_text (str): The full text of the original document.
        template_text (str): The template text used for context.
    
    Returns:
        str: A summary of the document.
    """
    
    if len(document_text) > LONG_DOC_THRESHOLD:
        document_text = summarize_long_document(document_text)

    # Proceed with the original summarization method using a custom prompt.
    system_prompt = summarize_document_prompt
    user_prompt = (
        f"Template:\n{template_text}\n\n"
        f"Document:\n{document_text}\n\n"
        "Summary (bullet points):"
    )
    messages = [
        {"role": "system", "content": system_prompt},
        {"role": "user", "content": user_prompt}
    ]
    try:
        response = client.chat.completions.create(
            messages=messages,
            model="gpt-4o",
            temperature=0.5,
            max_tokens=500
        )
        summary = response.choices[0].message.content
        return summary
    except Exception as e:
        print(f"Error summarizing document: {e}")
        return document_text

# =============================================================================
# Function: Generate Document (with Optional Context)
# =============================================================================
def generate_document(template_text, info_text, context_chunks=None):
    """
    Generate a document by combining the template and a summary of the original information text.
    If additional context chunks are provided, include them in the prompt.
    
    Args:
        template_text (str): The document template.
        info_text (str): The original document text.
        context_chunks (list, optional): A list of Document objects retrieved from additional context files.
    
    Returns:
        str: The generated document.
    """
    summarized_info = summarize_document(info_text, template_text)
    
    additional_context_text = ""
    if context_chunks:
        additional_context_text = "\n".join([doc.page_content for doc in context_chunks])
    
    system_prompt = generate_document_prompt  # Detailed instructions from prompts.json.
    user_prompt = f"Template:\n{template_text}\n\n"
    if additional_context_text:
        user_prompt += f"Additional Context:\n{additional_context_text}\n\n"
    user_prompt += f"Original Document Summary:\n{summarized_info}"
    
    messages = [
        {"role": "system", "content": system_prompt},
        {"role": "user", "content": user_prompt}
    ],
    try:
        response = client.chat.completions.create(
            messages=messages[0],
            model="gpt-4o",
            temperature=0.5,
            max_tokens=4096
        )
        generated_document = response.choices[0].message.content
    except Exception as e:
        generated_document = f"An error occurred while generating the document: {str(e)}"
    
    return generated_document

# =============================================================================
# Route: Index Page
# =============================================================================
@app.route('/')
def index():
    """
    Render the homepage with the file upload form.
    """
    return render_template('index.html')

# =============================================================================
# Route: Document Generation
# =============================================================================
@app.route('/generate', methods=['POST'])
def generate():
    """
    Handle file uploads, including optional additional context documents.
    Process and parse the files (supporting .txt, .md, .docx, .pdf),
    index and retrieve relevant chunks from context documents (if provided),
    summarize the original document with reference to the template,
    generate the final document, and render the result.
    """
    if 'template_file' not in request.files or 'info_file' not in request.files:
        flash("Both template and information files are required.")
        return redirect(request.url)

    template_file = request.files['template_file']
    info_file = request.files['info_file']

    if template_file.filename == "" or info_file.filename == "":
        flash("Please select both a template file and an information file.")
        return redirect(url_for('index'))

    try:
        template_text = read_uploaded_file(template_file)
        info_text = read_uploaded_file(info_file)
    except Exception as e:
        flash(f"Error reading files: {str(e)}")
        return redirect(url_for('index'))

    # Process additional context documents if provided.
    context_files = request.files.getlist('context_files')
    retrieved_docs = process_context_files(context_files, template_text)
    
    # Generate the final document using the provided files and any retrieved context.
    final_document = generate_document(template_text, info_text, context_chunks=retrieved_docs)
    return render_template('result.html', document=final_document)

# =============================================================================
# Main Entry Point
# =============================================================================
if __name__ == '__main__':
    app.run(debug=True)
