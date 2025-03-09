import os
from werkzeug.utils import secure_filename
from flask import current_app

def read_uploaded_file(uploaded_file):
    """
    Read the uploaded file and extract text depending on its file type.
    Supports .txt, .md, .docx, and .pdf files.
    
    Args:
        uploaded_file: The uploaded file object from request.files.
        
    Returns:
        str: The extracted text content from the file.
    """
    filename = secure_filename(uploaded_file.filename)
    ext = os.path.splitext(filename)[1].lower()
    
    # For text-based files
    if ext in ['.txt', '.md']:
        return uploaded_file.read().decode('utf-8')
    
    # For DOCX files
    elif ext == '.docx':
        from docx import Document
        doc = Document(uploaded_file)
        full_text = []
        for para in doc.paragraphs:
            full_text.append(para.text)
        return "\n".join(full_text)
    
    # For PDF files
    elif ext == '.pdf':
        import PyPDF2
        uploaded_file.seek(0)
        reader = PyPDF2.PdfReader(uploaded_file)
        pages_text = []
        for page in reader.pages:
            text = page.extract_text()
            if text:
                pages_text.append(text)
        return "\n".join(pages_text)
    
    # Fallback to reading as text
    else:
        return uploaded_file.read().decode('utf-8')

def read_file_content(file_object):
    """
    A simplified version of read_uploaded_file that doesn't reset file pointer.
    Useful for API where file might need to be processed differently.
    
    Args:
        file_object: The file object to read.
        
    Returns:
        str: The extracted text content.
    """
    return read_uploaded_file(file_object)

def process_context_files(context_files, query_text):
    """
    Process additional context files uploaded by the user.
    
    Args:
        context_files (list): List of file objects from the request.
        query_text (str): The text to use as a query for similarity search.
        
    Returns:
        list: A list of Document objects that are most relevant or an empty list if no context is provided.
    """
    # If no context files are provided, return an empty list
    if not context_files or context_files[0].filename == "":
        return []
    
    # Read and extract text from each context file
    context_texts = []
    for file in context_files:
        context_texts.append(read_uploaded_file(file))
    
    # Split the text into chunks and convert each chunk into a Document
    from langchain.text_splitter import RecursiveCharacterTextSplitter
    from langchain.docstore.document import Document
    all_context_docs = []
    text_splitter = RecursiveCharacterTextSplitter(chunk_size=1000, chunk_overlap=200)
    for doc_text in context_texts:
        chunks = text_splitter.split_text(doc_text)
        for chunk in chunks:
            all_context_docs.append(Document(page_content=chunk))
    
    # Build a vector store from the Document objects and perform similarity search
    from langchain.embeddings import OpenAIEmbeddings
    from langchain.vectorstores import FAISS
    
    # Get API key from Flask app config
    api_key = current_app.config.get('OPENAI_API_KEY')
    
    embeddings = OpenAIEmbeddings(openai_api_key=api_key)
    vector_store = FAISS.from_documents(all_context_docs, embeddings)
    retrieved_docs = vector_store.similarity_search(query_text, k=5)
    return retrieved_docs