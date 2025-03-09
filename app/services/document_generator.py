import io
from flask import current_app
from app.services.openai_service import generate_completion, get_prompts

def summarize_long_document(document_text):
    """
    Summarize a long document using LangChain's text splitter and summarization chain.
    
    Args:
        document_text (str): The document text to summarize.
        
    Returns:
        str: A summary of the document.
    """
    from langchain.text_splitter import RecursiveCharacterTextSplitter
    from langchain.docstore.document import Document
    
    # Split the long document into smaller chunks
    text_splitter = RecursiveCharacterTextSplitter(chunk_size=1000, chunk_overlap=200)
    texts = text_splitter.split_text(document_text)
    
    # Convert each chunk into a Document object
    docs = [Document(page_content=t) for t in texts]
    
    # Use the chat-based LLM wrapper for LangChain
    from langchain.chat_models import ChatOpenAI
    llm = ChatOpenAI(temperature=0.5, model="gpt-4o", openai_api_key=current_app.config.get('OPENAI_API_KEY'))
    
    # Load a summarization chain (map_reduce is a good choice for long documents)
    from langchain.chains.summarize import load_summarize_chain
    chain = load_summarize_chain(llm, chain_type="map_reduce")
    
    # Process each document chunk individually
    partial_summaries = []
    for doc in docs:
        summary = chain.run([doc])
        partial_summaries.append(summary)
    
    # Combine the partial summaries into a single summary
    long_summary = "\n".join(partial_summaries)
    return long_summary

def summarize_document(document_text, template_text, LONG_DOC_THRESHOLD=3000):
    """
    Summarize the given document text with reference to the template.
    
    Args:
        document_text (str): The full text of the original document.
        template_text (str): The template text used for context.
        LONG_DOC_THRESHOLD (int): Character threshold to determine if document is "long".
        
    Returns:
        str: A summary of the document.
    """
    if len(document_text) > LONG_DOC_THRESHOLD:
        document_text = summarize_long_document(document_text)

    # Get prompts from application configuration
    prompts = get_prompts()
    
    # Create prompt for summarization
    system_prompt = prompts.get("summarize_document_prompt")
    user_prompt = (
        f"Template:\n{template_text}\n\n"
        f"Document:\n{document_text}\n\n"
        "Summary (bullet points):"
    )
    messages = [
        {"role": "system", "content": system_prompt},
        {"role": "user", "content": user_prompt}
    ]
    
    # Generate summary using OpenAI
    try:
        summary = generate_completion(
            messages=messages,
            model="gpt-4o",
            temperature=0.5,
            max_tokens=500
        )
        return summary
    except Exception as e:
        current_app.logger.error(f"Error summarizing document: {e}")
        return document_text

def generate_document(template_text, info_text, context_chunks=None):
    """
    Generate a document by combining the template and a summary of the original information text.
    
    Args:
        template_text (str): The document template.
        info_text (str): The original document text.
        context_chunks (list, optional): A list of Document objects retrieved from additional context files.
        
    Returns:
        str: The generated document.
    """
    # Summarize the original document with reference to the template
    summarized_info = summarize_document(info_text, template_text)
    
    # Prepare additional context if available
    additional_context_text = ""
    if context_chunks:
        additional_context_text = "\n".join([doc.page_content for doc in context_chunks])
    
    # Get prompts from application configuration
    prompts = get_prompts()
    
    # Create prompt for document generation
    system_prompt = prompts.get("generate_document_from_template_prompt")
    user_prompt = f"Template:\n{template_text}\n\n"
    if additional_context_text:
        user_prompt += f"Additional Context:\n{additional_context_text}\n\n"
    user_prompt += f"Original Document Summary:\n{summarized_info}"
    
    messages = [
        {"role": "system", "content": system_prompt},
        {"role": "user", "content": user_prompt}
    ]
    
    # Generate document using OpenAI
    try:
        generated_document = generate_completion(
            messages=messages,
            model="gpt-4o",
            temperature=0.5,
            max_tokens=4096
        )
        return generated_document
    except Exception as e:
        current_app.logger.error(f"Error generating document: {str(e)}")
        return f"An error occurred while generating the document: {str(e)}"

def generate_docx(text):
    """
    Generate a DOCX file from the given text.
    
    Args:
        text (str): The text to include in the document.
        
    Returns:
        BytesIO: An in-memory DOCX file.
    """
    from docx import Document as DocxDocument
    doc = DocxDocument()
    # Add a title and paragraph; customize formatting as needed
    doc.add_heading("Generated Document", level=1)
    doc.add_paragraph(text)
    f = io.BytesIO()
    doc.save(f)
    f.seek(0)
    return f

def generate_pdf(text):
    """
    Generate a PDF file from the given text.
    
    Args:
        text (str): The text to include in the document.
        
    Returns:
        BytesIO: An in-memory PDF file.
    """
    from weasyprint import HTML
    # Wrap the text in basic HTML formatting
    html_str = f"<html><body><h1>Generated Document</h1><pre>{text}</pre></body></html>"
    pdf = HTML(string=html_str).write_pdf()
    f = io.BytesIO(pdf)
    f.seek(0)
    return f